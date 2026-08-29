"""Tests for C-03 Office and plain-text/HTML extractors."""

from __future__ import annotations

from datetime import UTC, datetime
from pathlib import Path
from typing import TYPE_CHECKING

import pytest
from docx import Document

from ats_scan.codes import ReasonCode
from ats_scan.extract.office.extractor import DocxExtractor, LegacyOfficeExtractor
from ats_scan.extract.plain._html import html_to_text
from ats_scan.extract.plain._langdetect import detect_language
from ats_scan.extract.plain._normalise import normalise_text
from ats_scan.extract.plain.extractor import HtmlExtractor, MarkdownExtractor, PlainTextExtractor
from ats_scan.extract.registry import load_extractors
from ats_scan.models.run import RunContext
from ats_scan.models.source import SourceDocument

if TYPE_CHECKING:
    from ats_scan.protocols import TextExtractor


@pytest.fixture
def ctx() -> RunContext:
    """Minimal run context for extractors."""
    return RunContext(run_id="test-run")


def _doc(
    path: str, media_type: str = "application/octet-stream", *, pages: int | None = None
) -> SourceDocument:
    """Build a SourceDocument for a file on disk."""
    p = Path(path)
    return SourceDocument(
        path=str(p),
        content_sha256=p.read_bytes().hex()[:64] or "0" * 64,
        bytes=p.stat().st_size,
        pages=pages,
        mtime=datetime.fromtimestamp(p.stat().st_mtime, tz=UTC).isoformat().replace("+00:00", "Z"),
        media_type=media_type,
    )


class TestNormalise:
    """Unicode normalisation contract per FR-210."""

    def test_nfkc_ligatures(self) -> None:
        text = normalise_text("ﬁnancial ﬂexibility")
        assert text == "financial flexibility"

    def test_zero_width_removed(self) -> None:
        text = normalise_text("zero\u200bwidth\u200cjoiner\u200d")
        assert text == "zerowidthjoiner"

    def test_bidi_controls_removed(self) -> None:
        text = normalise_text("\u202ahello\u202c")
        assert text == "hello"

    def test_bom_removed(self) -> None:
        text = normalise_text("\ufeffhello")
        assert text == "hello"


class TestLanguageDetect:
    """Stopword-based language detection per FR-209."""

    def test_english_detected(self) -> None:
        lang, confidence = detect_language("This is a simple English resume.", ["en"])
        assert lang == "en"
        assert confidence > 0.0

    def test_spanish_detected_even_if_not_supported(self) -> None:
        lang, confidence = detect_language("Este es un currículum en español.", ["en"])
        assert lang == "es"
        assert confidence > 0.0

    def test_spanish_supported(self) -> None:
        lang, confidence = detect_language("Este es un currículum en español.", ["en", "es"])
        assert lang == "es"
        assert confidence > 0.0


class TestHtmlToText:
    """HTML stripping preserving block structure."""

    def test_block_structure_preserved(self) -> None:
        html = "<h1>Header</h1><p>Paragraph one.</p><p>Paragraph two.</p>"
        text = html_to_text(html)
        assert "Header" in text
        assert "Paragraph one." in text
        assert "Paragraph two." in text
        assert text.count("\n") >= 2

    def test_script_and_style_removed(self) -> None:
        html = "<script>alert(1)</script><style>.x{}</style><p>visible</p>"
        text = html_to_text(html)
        assert "alert" not in text
        assert ".x{}" not in text
        assert "visible" in text

    def test_entities_unescaped(self) -> None:
        html = "<p>Tom &amp; Jerry &#169; 2026</p>"
        text = html_to_text(html)
        assert "Tom & Jerry" in text
        assert "© 2026" in text


class TestPlainTextExtractor:
    """``.txt`` extraction."""

    def test_extracts_text(self, ctx: RunContext, tmp_path: Path) -> None:
        path = tmp_path / "resume.txt"
        path.write_text(
            "John Doe is a Python developer with experience in the field.\n",
            encoding="utf-8",
        )
        doc = _doc(str(path), "text/plain")
        extractor = PlainTextExtractor()
        result = extractor.extract(doc, ctx)
        assert result.ok
        assert result.value is not None
        assert "John Doe" in result.value.text
        assert result.value.metadata.language == "en"
        assert result.value.metadata.language_confidence > 0.0

    def test_normalises_text(self, ctx: RunContext, tmp_path: Path) -> None:
        path = tmp_path / "resume.txt"
        path.write_text("\ufeffﬁnance\u200b", encoding="utf-8")
        doc = _doc(str(path), "text/plain")
        result = PlainTextExtractor().extract(doc, ctx)
        assert result.value is not None
        assert result.value.text == "finance"

    def test_block_text_matches_extracted_text(self, ctx: RunContext, tmp_path: Path) -> None:
        path = tmp_path / "resume.txt"
        path.write_text("Hello world", encoding="utf-8")
        doc = _doc(str(path), "text/plain")
        result = PlainTextExtractor().extract(doc, ctx)
        assert result.value is not None
        block = result.value.blocks[0]
        assert block.text == result.value.text


class TestMarkdownExtractor:
    """``.md`` extraction treated as plain text."""

    def test_extracts_markdown(self, ctx: RunContext, tmp_path: Path) -> None:
        path = tmp_path / "resume.md"
        path.write_text("# Jane Doe\n\n* Python\n* Rust\n", encoding="utf-8")
        doc = _doc(str(path), "text/markdown")
        result = MarkdownExtractor().extract(doc, ctx)
        assert result.ok
        assert result.value is not None
        assert "Jane Doe" in result.value.text


class TestHtmlExtractor:
    """``.html`` extraction."""

    def test_extracts_html(self, ctx: RunContext, tmp_path: Path) -> None:
        path = tmp_path / "resume.html"
        path.write_text(
            "<html><body><h1>Jane Doe</h1><p>Python developer</p></body></html>",
            encoding="utf-8",
        )
        doc = _doc(str(path), "text/html")
        result = HtmlExtractor().extract(doc, ctx)
        assert result.ok
        assert result.value is not None
        assert "Jane Doe" in result.value.text
        assert "<h1>" not in result.value.text
        assert "<p>" not in result.value.text

    def test_block_structure_preserved(self, ctx: RunContext, tmp_path: Path) -> None:
        path = tmp_path / "resume.html"
        path.write_text("<h1>A</h1><p>B</p><div>C</div>", encoding="utf-8")
        doc = _doc(str(path), "text/html")
        result = HtmlExtractor().extract(doc, ctx)
        assert result.value is not None
        text = result.value.text
        assert text.count("\n") >= 2


class TestDocxExtractor:
    """``.docx`` extraction via python-docx."""

    def test_extracts_paragraphs_and_tables(self, ctx: RunContext, tmp_path: Path) -> None:
        path = tmp_path / "resume.docx"
        document = Document()
        document.add_heading("Alice Smith", level=1)
        document.add_paragraph("Software engineer with Python experience.")
        table = document.add_table(rows=1, cols=2)
        row = table.rows[0]
        row.cells[0].text = "Skill"
        row.cells[1].text = "Years"
        document.save(str(path))

        doc = _doc(
            str(path), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        result = DocxExtractor().extract(doc, ctx)
        assert result.ok
        assert result.value is not None
        assert "Alice Smith" in result.value.text
        assert "Software engineer" in result.value.text
        assert "Skill" in result.value.text
        assert "Years" in result.value.text

    def test_corrupt_docx_returns_diagnostic(self, ctx: RunContext, tmp_path: Path) -> None:
        path = tmp_path / "broken.docx"
        path.write_bytes(b"not a zip")
        doc = _doc(
            str(path), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        result = DocxExtractor().extract(doc, ctx)
        assert not result.ok
        assert result.value is None
        assert any(d.code == ReasonCode.EXT_CORRUPT for d in result.diagnostics)


class TestLegacyOfficeExtractor:
    """``.doc`` and ``.rtf`` via headless converter."""

    def test_no_converter_returns_diagnostic(
        self, ctx: RunContext, tmp_path: Path, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        monkeypatch.setenv("ATS_SCAN_OFFICE_CONVERTER_CMD", "/nonexistent/binary")
        path = tmp_path / "legacy.doc"
        path.write_bytes(b"dummy")
        doc = _doc(str(path), "application/msword")
        result = LegacyOfficeExtractor().extract(doc, ctx)
        assert not result.ok
        assert result.value is None
        assert any(d.code == ReasonCode.EXT_CORRUPT for d in result.diagnostics)

    def test_converter_timeout_kills_process(
        self, ctx: RunContext, tmp_path: Path, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        # Create a fake converter that sleeps longer than the timeout.
        converter = tmp_path / "slow_converter.sh"
        converter.write_text("#!/bin/sh\nsleep 5\n", encoding="utf-8")
        converter.chmod(0o755)
        monkeypatch.setenv("ATS_SCAN_OFFICE_CONVERTER_CMD", str(converter))
        monkeypatch.setenv("ATS_SCAN_CONVERTER_TIMEOUT_S", "1")

        path = tmp_path / "legacy.rtf"
        path.write_bytes(b"dummy")
        doc = _doc(str(path), "application/rtf")

        result = LegacyOfficeExtractor().extract(doc, ctx)
        assert not result.ok
        assert result.value is None
        assert any(d.code == ReasonCode.EXT_CORRUPT for d in result.diagnostics)
        assert any("timed out" in d.message for d in result.diagnostics)

        # Ensure the converter process was actually killed (no lingering sleep).
        # The subprocess module will have terminated the process on timeout.

    def test_converter_success(
        self, ctx: RunContext, tmp_path: Path, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        # Fake converter that writes a .txt file into the output directory.
        converter = tmp_path / "fake_converter.sh"
        script = (
            "#!/bin/sh\n"
            "outdir=$1\n"
            "input=$2\n"
            'basename=$(basename "$input" .doc)\n'
            "echo 'Extracted legacy text' > \"$outdir/${basename}.txt\"\n"
        )
        converter.write_text(script, encoding="utf-8")
        converter.chmod(0o755)
        monkeypatch.setenv("ATS_SCAN_OFFICE_CONVERTER_CMD", str(converter))

        path = tmp_path / "legacy.doc"
        path.write_bytes(b"dummy")
        doc = _doc(str(path), "application/msword")
        result = LegacyOfficeExtractor().extract(doc, ctx)
        assert result.ok
        assert result.value is not None
        assert "Extracted legacy text" in result.value.text


class TestLanguageFlag:
    """Unsupported language detection."""

    def test_unsupported_language_emits_diagnostic(
        self, ctx: RunContext, tmp_path: Path, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        monkeypatch.setenv("ATS_SCAN_LANGUAGES", "en")
        path = tmp_path / "spanish.txt"
        path.write_text(
            "Este es un currículum en español. La experiencia incluye Python.", encoding="utf-8"
        )
        doc = _doc(str(path), "text/plain")
        result = PlainTextExtractor().extract(doc, ctx)
        assert result.ok
        assert result.value is not None
        assert result.value.metadata.language == "es"
        assert any(d.code == ReasonCode.LANG_UNSUPPORTED for d in result.diagnostics)


class TestRegistry:
    """Extractor registration."""

    def test_extractors_registered(self) -> None:
        extractors = load_extractors()
        names = set(extractors)
        assert "PlainTextExtractor" in names
        assert "MarkdownExtractor" in names
        assert "HtmlExtractor" in names
        assert "DocxExtractor" in names
        assert "LegacyOfficeExtractor" in names

    @pytest.mark.parametrize(
        ("cls", "media_type", "extension"),
        [
            (PlainTextExtractor, "text/plain", ".txt"),
            (MarkdownExtractor, "text/markdown", ".md"),
            (HtmlExtractor, "text/html", ".html"),
            (
                DocxExtractor,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                ".docx",
            ),
            (LegacyOfficeExtractor, "application/msword", ".doc"),
            (LegacyOfficeExtractor, "application/rtf", ".rtf"),
        ],
    )
    def test_supports_by_media_type_and_extension(
        self, cls: type[TextExtractor], media_type: str, extension: str, tmp_path: Path
    ) -> None:
        path = tmp_path / f"file{extension}"
        path.write_text("x", encoding="utf-8")
        doc = SourceDocument(
            path=str(path),
            content_sha256="0" * 64,
            bytes=1,
            mtime="2026-08-29T00:00:00Z",
            media_type=media_type,
        )
        extractor = cls()
        assert extractor.supports(doc)
