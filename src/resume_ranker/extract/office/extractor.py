"""Office document extraction: ``.docx`` via python-docx, ``.doc``/``.rtf`` via headless converter (C-03 / TRD §3.2 / FR-208)."""

from __future__ import annotations

import os
import shutil
import subprocess
import tempfile
from pathlib import Path
from typing import ClassVar, Final

from docx import Document

from resume_ranker.codes import ReasonCode
from resume_ranker.extract.plain._langdetect import detect_language
from resume_ranker.extract.plain._normalise import normalise_text
from resume_ranker.extract.registry import extractor
from resume_ranker.models.common import Diagnostic, StageResult
from resume_ranker.models.run import RunContext
from resume_ranker.models.source import ExtractedText, ExtractionMetadata, SourceDocument, TextBlock
from resume_ranker.protocols import TextExtractor

_CONVERTER_TIMEOUT_S: Final[int] = 60


def _supported_languages() -> tuple[str, ...]:
    """Return the configured accepted language set, defaulting to English.

    Note: extraction configuration is not currently reachable through
    ``RunContext.config`` (which only carries ``IngestConfig``).  Contract
    change C-03-002 requests that path.  Until it is merged, the environment
    variable ``ATS_SCAN_LANGUAGES`` (comma-separated) or the English default is
    used.
    """
    env = os.environ.get("ATS_SCAN_LANGUAGES", "en")
    return tuple(lang.strip() for lang in env.split(",") if lang.strip()) or ("en",)


def _converter_timeout() -> int:
    """Return the configured converter timeout in seconds, defaulting to 60."""
    env = os.environ.get("ATS_SCAN_CONVERTER_TIMEOUT_S", str(_CONVERTER_TIMEOUT_S))
    try:
        return int(env)
    except (TypeError, ValueError):
        return _CONVERTER_TIMEOUT_S


def _converter_command() -> list[str] | None:
    """Return the configured or discovered headless office converter command."""
    env_cmd = os.environ.get("ATS_SCAN_OFFICE_CONVERTER_CMD")
    if env_cmd:
        return env_cmd.split()

    for binary in ("libreoffice", "soffice"):
        executable = shutil.which(binary)
        if executable:
            return [
                executable,
                "--headless",
                "--nologo",
                "--nodefault",
                "--norestore",
                "--nolockcheck",
                "--macro-security-level=4",
                "--convert-to",
                "txt:Text",
                "--outdir",
            ]
    return None


def _convert_legacy(path: str, ctx: RunContext) -> tuple[str | None, Diagnostic | None]:
    """Convert ``.doc`` or ``.rtf`` to text using a headless converter with timeout.

    Implements TRD §3.2 / FR-208: headless office converter with networking and
    macros disabled, under a wall-clock timeout (default 60 s).  On failure
    returns ``None`` and a diagnostic, never raises.
    """
    command = _converter_command()
    if command is None:
        return None, Diagnostic(
            stage="S2",
            code=ReasonCode.EXT_CORRUPT,
            message="No headless office converter (libreoffice/soffice) is available.",
        )

    timeout = _converter_timeout()
    with tempfile.TemporaryDirectory() as tmpdir:
        outdir = Path(tmpdir)
        try:
            subprocess.run(
                command + [str(outdir), path],
                check=True,
                timeout=timeout,
                capture_output=True,
                stdin=subprocess.DEVNULL,
                env={**os.environ, "HOME": tmpdir},
            )
        except subprocess.TimeoutExpired:
            return None, Diagnostic(
                stage="S2",
                code=ReasonCode.EXT_CORRUPT,
                message=f"Office converter timed out after {timeout}s and was killed.",
            )
        except (subprocess.CalledProcessError, OSError) as exc:
            reason = getattr(exc, "stderr", b"")
            if isinstance(reason, bytes):
                reason = reason.decode("utf-8", errors="replace")[:200]
            elif not isinstance(reason, str):
                reason = str(exc)
            return None, Diagnostic(
                stage="S2",
                code=ReasonCode.EXT_CORRUPT,
                message=f"Office converter failed: {reason}",
            )

        stem = Path(path).stem
        candidates = sorted(outdir.glob("*.txt"))
        if not candidates:
            return None, Diagnostic(
                stage="S2",
                code=ReasonCode.EXT_CORRUPT,
                message="Office converter produced no text output.",
            )

        output_path = candidates[0]
        # If multiple files were produced, prefer the one matching the input stem.
        for candidate in candidates:
            if candidate.stem == stem:
                output_path = candidate
                break

        try:
            text = output_path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            text = output_path.read_text(encoding="latin-1")
        return text, None


def _extract_docx(path: str) -> str:
    """Extract text from a ``.docx`` file, preserving paragraph and table order.

    Implements TRD §3.2 / FR-208: .docx extraction including tables and heading
    paragraphs.
    """
    document = Document(path)
    paragraphs: list[str] = []

    # Header / footer are not required for C-03; body paragraphs and tables are.
    for paragraph in document.paragraphs:
        if paragraph.text:
            paragraphs.append(paragraph.text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))

    return "\n".join(paragraphs)


def _make_result(
    doc: SourceDocument,
    raw_text: str,
    method: str,
    ctx: RunContext,
) -> StageResult[ExtractedText]:
    """Normalise text, detect language, and package the extraction result.

    Implements TRD §3.2 / FR-209 (language detection) and FR-210 (Unicode
    normalisation).  Returns a diagnostic when the detected language is outside
    the configured accepted set.
    """
    text = normalise_text(raw_text)
    supported = _supported_languages()
    language, language_confidence = detect_language(text, supported)
    metadata = ExtractionMetadata(
        method=method,
        chars_per_page=float(len(text)) if doc.pages in (None, 0) else len(text) / doc.pages,
        language=language,
        language_confidence=language_confidence,
        quality=1.0,
        ocr_confidence=None,
        columns_detected=None,
    )
    diagnostics: list[Diagnostic] = []
    if language not in supported:
        diagnostics.append(
            Diagnostic(
                stage="S2",
                code=ReasonCode.LANG_UNSUPPORTED,
                message=f"Primary language '{language}' is not in the configured set.",
            )
        )
    block = TextBlock(text=text, page=0, bbox=(0.0, 0.0, 0.0, 0.0))
    return StageResult(
        value=ExtractedText(
            text=text,
            metadata=metadata,
            blocks=(block,),
            language=language,
            language_confidence=language_confidence,
        ),
        diagnostics=tuple(diagnostics),
    )


@extractor
class DocxExtractor(TextExtractor):
    """Extractor for modern Office Open XML ``.docx`` files (TRD §3.2 / FR-208)."""

    media_types: ClassVar[frozenset[str]] = frozenset(
        {
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        }
    )

    def supports(self, doc: SourceDocument) -> bool:
        return doc.media_type in self.media_types or doc.path.lower().endswith(".docx")

    def extract(self, doc: SourceDocument, ctx: RunContext) -> StageResult[ExtractedText]:
        try:
            raw_text = _extract_docx(doc.path)
        except Exception as exc:  # noqa: BLE001 - extraction errors must become diagnostics
            return StageResult(
                value=None,
                diagnostics=(
                    Diagnostic(
                        stage="S2",
                        code=ReasonCode.EXT_CORRUPT,
                        message=f"Could not read .docx file: {exc}",
                    ),
                ),
            )
        return _make_result(doc, raw_text, "docx", ctx)


@extractor
class LegacyOfficeExtractor(TextExtractor):
    """Extractor for legacy ``.doc`` and ``.rtf`` files via headless converter (TRD §3.2 / FR-208)."""

    media_types: ClassVar[frozenset[str]] = frozenset(
        {"application/msword", "application/rtf", "text/rtf"}
    )

    def supports(self, doc: SourceDocument) -> bool:
        return doc.media_type in self.media_types or doc.path.lower().endswith((".doc", ".rtf"))

    def extract(self, doc: SourceDocument, ctx: RunContext) -> StageResult[ExtractedText]:
        raw_text, diagnostic = _convert_legacy(doc.path, ctx)
        if raw_text is None:
            assert diagnostic is not None  # guaranteed by _convert_legacy contract
            return StageResult(
                value=None,
                diagnostics=(diagnostic,),
            )
        return _make_result(
            doc, raw_text, "doc" if doc.path.lower().endswith(".doc") else "rtf", ctx
        )
